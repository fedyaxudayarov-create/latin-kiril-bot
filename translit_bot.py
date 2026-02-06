import os
import re
import uuid
import logging
from typing import Tuple

from telegram import (
    Update,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
)
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    CallbackQueryHandler,
    ContextTypes,
    filters,
)

from docx import Document as DocxDocument
from openpyxl import load_workbook
import fitz  # PyMuPDF
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# =========================
# SOZLAMALAR
# =========================
BOT_TOKEN = "8334778816:AAGjZ2lC0VUOICgZGjQRqjMB7Phaf2t54bM"

WORKDIR = "work"
os.makedirs(WORKDIR, exist_ok=True)

logging.basicConfig(level=logging.INFO)


# =========================
# LATIN <-> KIRILL
# =========================
_APOS = ["'", "’", "`", "ʼ", "ʻ", "‘"]

def _norm_apos(s: str) -> str:
    for a in _APOS[1:]:
        s = s.replace(a, "'")
    return s

_LAT2CYR_MULTI = [
    ("o'", "ў"), ("g'", "ғ"),
    ("sh", "ш"), ("ch", "ч"), ("ng", "нг"),
    ("ya", "я"), ("yo", "ё"), ("yu", "ю"),
]

_LAT2CYR_SINGLE = {
    "a": "а", "b": "б", "d": "д", "e": "е", "f": "ф", "g": "г",
    "h": "ҳ", "i": "и", "j": "ж", "k": "к", "l": "л", "m": "м",
    "n": "н", "o": "о", "p": "п", "q": "қ", "r": "р", "s": "с",
    "t": "т", "u": "у", "v": "в", "x": "х", "y": "й", "z": "з",
    "'": "ъ",
}

_CYR2LAT_MULTI = [
    ("нг", "ng"), ("ш", "sh"), ("ч", "ch"),
    ("ў", "o'"), ("ғ", "g'"),
    ("я", "ya"), ("ё", "yo"), ("ю", "yu"),
]

_CYR2LAT_SINGLE = {
    "а": "a", "б": "b", "д": "d", "е": "e", "ф": "f", "г": "g",
    "ҳ": "h", "и": "i", "ж": "j", "к": "k", "л": "l", "м": "m",
    "н": "n", "о": "o", "п": "p", "қ": "q", "р": "r", "с": "s",
    "т": "t", "у": "u", "в": "v", "х": "x", "й": "y", "з": "z",
    "ъ": "'", "ь": "",
}

def latin_to_cyrillic(text: str) -> str:
    text = _norm_apos(text)
    for src, dst in _LAT2CYR_MULTI:
        text = re.sub(src, dst, text, flags=re.IGNORECASE)

    out = []
    for ch in text:
        low = ch.lower()
        if low in _LAT2CYR_SINGLE:
            out.append(_LAT2CYR_SINGLE[low].upper() if ch.isupper() else _LAT2CYR_SINGLE[low])
        else:
            out.append(ch)
    return "".join(out)

def cyrillic_to_latin(text: str) -> str:
    for src, dst in _CYR2LAT_MULTI:
        text = text.replace(src, dst).replace(src.upper(), dst.upper())
    out = []
    for ch in text:
        low = ch.lower()
        if low in _CYR2LAT_SINGLE:
            out.append(_CYR2LAT_SINGLE[low].upper() if ch.isupper() else _CYR2LAT_SINGLE[low])
        else:
            out.append(ch)
    return "".join(out)

def translit(text: str, direction: str) -> str:
    return latin_to_cyrillic(text) if direction == "lat2cyr" else cyrillic_to_latin(text)


# =========================
# FILE ISHLOV
# =========================
def process_docx(inp, outp, d):
    doc = DocxDocument(inp)
    for p in doc.paragraphs:
        p.text = translit(p.text, d)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    p.text = translit(p.text, d)
    doc.save(outp)

def process_xlsx(inp, outp, d):
    wb = load_workbook(inp)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    cell.value = translit(cell.value, d)
    wb.save(outp)

def process_pdf(inp, outp, d) -> Tuple[bool, str]:
    doc = fitz.open(inp)
    text = "\n".join(p.get_text() for p in doc)
    doc.close()
    if not text.strip():
        return False, "PDF матнли эмас"
    text = translit(text, d)

    font = r"C:\Windows\Fonts\arial.ttf"
    c = canvas.Canvas(outp, pagesize=A4)
    if os.path.exists(font):
        pdfmetrics.registerFont(TTFont("A", font))
        c.setFont("A", 11)
    y = 800
    for line in text.splitlines():
        if y < 40:
            c.showPage()
            y = 800
        c.drawString(40, y, line)
        y -= 14
    c.save()
    return True, "OK"


# =========================
# TELEGRAM BOT
# =========================
def kb():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("Лотин → Кирилл", callback_data="lat2cyr")],
        [InlineKeyboardButton("Кирилл → Лотин", callback_data="cyr2lat")]
    ])

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Файл ташланг (DOCX / XLSX / PDF)")

async def on_doc(update: Update, context: ContextTypes.DEFAULT_TYPE):
    doc = update.message.document
    ext = os.path.splitext(doc.file_name)[1].lower()
    if ext not in [".docx", ".xlsx", ".pdf"]:
        await update.message.reply_text("Фақат DOCX, XLSX, PDF")
        return
    f = await context.bot.get_file(doc.file_id)
    path = os.path.join(WORKDIR, str(uuid.uuid4()) + ext)
    await f.download_to_drive(path)
    context.user_data["file"] = (path, ext, doc.file_name)
    await update.message.reply_text("Қайсига ўгирилади?", reply_markup=kb())

async def on_choice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query
    await q.answer()
    d = q.data
    path, ext, name = context.user_data["file"]
    out = os.path.join(WORKDIR, str(uuid.uuid4()) + ext)

    if ext == ".docx":
        process_docx(path, out, d)
    elif ext == ".xlsx":
        process_xlsx(path, out, d)
    else:
        ok, msg = process_pdf(path, out, d)
        if not ok:
            await q.edit_message_text(msg)
            return

    with open(out, "rb") as f:
        await context.bot.send_document(q.message.chat_id, f, filename=name)

def main():
    application = Application.builder().token(BOT_TOKEN).build()

    application.add_handler(CommandHandler("start", start))
    application.add_handler(MessageHandler(filters.Document.ALL, on_doc))
    application.add_handler(CallbackQueryHandler(on_choice))

    print("Bot ishga tushdi...")
    application.run_polling()


if __name__ == "__main__":
    main()
